from docx import Document
from docx.shared import Pt, Cm
import re

md_path = 'Project_Report.md'
docx_path = 'Project_Report.docx'

def add_paragraph_from_text(doc, text):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.name = 'Times New Roman'
    run.font.size = Pt(12)

with open(md_path, 'r', encoding='utf-8') as f:
    lines = f.readlines()

# create document and set margins
doc = Document()
sections = doc.sections
for section in sections:
    section.top_margin = Cm(2)
    section.bottom_margin = Cm(2)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2)

current_para = None
for line in lines:
    line = line.rstrip('\n')
    if line.startswith('# '):
        h = doc.add_heading(line[2:].strip(), level=1)
        h.runs[0].font.name = 'Times New Roman'
        h.runs[0].font.size = Pt(14)
    elif line.startswith('## '):
        h = doc.add_heading(line[3:].strip(), level=2)
        h.runs[0].font.name = 'Times New Roman'
        h.runs[0].font.size = Pt(14)
    elif line.startswith('### '):
        h = doc.add_heading(line[4:].strip(), level=3)
        h.runs[0].font.name = 'Times New Roman'
        h.runs[0].font.size = Pt(13)
    elif re.match(r'^!\[.*\]\(.*\)$', line):
        # image markdown: ![alt](path)
        m = re.match(r'^!\[(.*)\]\((.*)\)$', line)
        if m:
            alt, path = m.groups()
            try:
                doc.add_picture(path, width=Cm(8))
                last_para = doc.paragraphs[-1]
                last_para.alignment = 0
            except Exception:
                add_paragraph_from_text(doc, f'[Image: {path}]')
    elif line.strip() == '':
        # blank line -> paragraph break
        current_para = None
    else:
        # normal paragraph
        if current_para is None:
            current_para = doc.add_paragraph()
            run = current_para.add_run(line)
            run.font.name = 'Times New Roman'
            run.font.size = Pt(12)
        else:
            current_para.add_run('\n' + line)

# Save
doc.save(docx_path)
print('Saved', docx_path)
